import pytest
import os
from pathlib import Path
import pandas as pd
import docx

@pytest.fixture
def sample_md_file(tmp_path):
    """Creates a sample markdown file."""
    content = """# Title
<!-- page_number: 1 -->
Intro text.

## Section 1
<!-- page_number: 2 -->
Section text.
"""
    p = tmp_path / "test.md"
    p.write_text(content, encoding='utf-8')
    return p

@pytest.fixture
def sample_docx_file(tmp_path):
    """Creates a sample docx file with text and table."""
    doc = docx.Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('Paragraph 1')
    
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header1"
    table.cell(0, 1).text = "Header2"
    table.cell(1, 0).text = "Val1"
    table.cell(1, 1).text = "Val2"
    
    p = tmp_path / "test.docx"
    doc.save(str(p))
    return p

@pytest.fixture
def sample_excel_file(tmp_path):
    """Creates a sample excel file with two sheets."""
    df1 = pd.DataFrame({'A': [1, 2], 'B': [3, 4]})
    df2 = pd.DataFrame({'C': [5, 6], 'D': [7, 8]})
    
    p = tmp_path / "test.xlsx"
    with pd.ExcelWriter(str(p)) as writer:
        df1.to_excel(writer, sheet_name='Sheet1', index=False)
        df2.to_excel(writer, sheet_name='Sheet2', index=False)
    return p
