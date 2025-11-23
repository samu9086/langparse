import pandas as pd
import docx
import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from langparse.autoparser import AutoParser
from langparse.chunkers.semantic import SemanticChunker
import os

def create_test_files():
    print("Creating test files...")
    
    # 1. Create Excel File
    df1 = pd.DataFrame({
        'ID': [1, 2, 3],
        'Name': ['Alice', 'Bob', 'Charlie'],
        'Role': ['Engineer', 'Designer', 'Manager']
    })
    df2 = pd.DataFrame({
        'Project': ['Alpha', 'Beta'],
        'Status': ['Done', 'In Progress']
    })
    
    with pd.ExcelWriter('test_data.xlsx') as writer:
        df1.to_excel(writer, sheet_name='Employees', index=False)
        df2.to_excel(writer, sheet_name='Projects', index=False)
        
    # 2. Create DOCX File
    doc = docx.Document()
    doc.add_heading('Project Report', 0)
    doc.add_paragraph('This is a test document created automatically.')
    
    doc.add_heading('Team Members', level=1)
    doc.add_paragraph('Here is the team list:')
    
    # Add a table
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Role'
    hdr_cells[2].text = 'Location'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Alice'
    row_cells[1].text = 'Dev'
    row_cells[2].text = 'NY'
    
    doc.save('test_report.docx')
    print("Files created: test_data.xlsx, test_report.docx")

def test_parsing():
    print("\n--- Testing Parsing ---")
    
    # Test Excel
    print("\n[Parsing Excel]")
    doc_xls = AutoParser.parse('test_data.xlsx')
    print(f"Content Preview:\n{doc_xls.content[:300]}...")
    
    # Test DOCX
    print("\n[Parsing DOCX]")
    doc_docx = AutoParser.parse('test_report.docx')
    print(f"Content Preview:\n{doc_docx.content}...")
    
    # Test Chunking (to verify page numbers)
    print("\n[Testing Chunking on Excel]")
    chunker = SemanticChunker()
    chunks = chunker.chunk(doc_xls)
    for chunk in chunks:
        print(f"Chunk Header: {chunk.metadata.get('header')} | Pages: {chunk.metadata.get('page_numbers')}")

if __name__ == "__main__":
    # Install dependencies if missing
    try:
        import pandas
        import openpyxl
        import docx
    except ImportError:
        print("Installing dependencies...")
        os.system("pip install pandas openpyxl python-docx")
        
    create_test_files()
    test_parsing()
    
    # Cleanup
    if os.path.exists("test_data.xlsx"):
        os.remove("test_data.xlsx")
    if os.path.exists("test_report.docx"):
        os.remove("test_report.docx")
